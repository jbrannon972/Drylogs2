#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys

def convert_md_to_docx(md_file, docx_file):
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create a new Document
    doc = Document()

    # Split content into lines
    lines = content.split('\n')

    i = 0
    in_code_block = False
    code_block_lines = []
    in_list = False

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle headings
        if line.startswith('#'):
            hash_count = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()

            if hash_count == 1:
                doc.add_heading(heading_text, level=1)
            elif hash_count == 2:
                doc.add_heading(heading_text, level=2)
            elif hash_count == 3:
                doc.add_heading(heading_text, level=3)
            elif hash_count == 4:
                doc.add_heading(heading_text, level=4)
            else:
                doc.add_heading(heading_text, level=5)

        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(text, style='List Number')

        # Handle horizontal rules
        elif line.strip() == '---' or line.strip() == '***':
            doc.add_paragraph('_' * 50)

        # Handle blockquotes
        elif line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)

        # Handle empty lines
        elif line.strip() == '':
            # Skip multiple empty lines
            if i > 0 and lines[i-1].strip() != '':
                doc.add_paragraph('')

        # Handle regular paragraphs
        else:
            # Process inline formatting
            text = line

            # Convert bold
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'__(.+?)__', r'\1', text)

            # Convert italic
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'_(.+?)_', r'\1', text)

            # Convert inline code
            text = re.sub(r'`(.+?)`', r'\1', text)

            if text.strip():
                doc.add_paragraph(text)

        i += 1

    # Save the document
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py input.md output.docx")
        sys.exit(1)

    convert_md_to_docx(sys.argv[1], sys.argv[2])
